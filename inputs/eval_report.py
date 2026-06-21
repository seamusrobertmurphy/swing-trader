"""
AA-evals: a shared writer for Chapter Three evaluation records.

Each evaluation run lands under outputs/AA-evals/ as a dated record, the way
AA-journal collects signal study sheets. Per-run files go in a date-named folder
(YYYY-MM-DD), and one consolidated table, evaluation-scores.md, sits at the root.

  AA-evals/
    evaluation-scores.md                         the consolidated index, one row per run
    2026-06-20/
      eval-head-to-head-20260620.md              readable record of truth
      eval-head-to-head-20260620.html            same plus charts embedded (self-contained)
      eval-head-to-head-20260620-*.png           charts

Each record reports Keller Metric 1 (precision/recall at the 60/40 trading threshold)
and Metric 3 (AUC split by volatility regime). Charts use matplotlib (no Chrome).
"""

import base64
import os
from datetime import datetime, timezone

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from sklearn.metrics import roc_auc_score, roc_curve

NAVY, GREEN, RED, GREY = "#0B3D66", "#2E8B57", "#B22222", "#888888"
PALETTE = [NAVY, GREEN, "#B8860B", RED, "#6A5ACD"]

INDEX_FILE = "evaluation-scores.md"
_INDEX_HEADERS = ["date", "evaluation type", "dataset", "best model", "test AUC",
                  "Best Model Precision", "Always Buys Precision", "Precision Change (%)",
                  "Net P&L/trade", "trades", "verdict", "record"]

GLOSSARY = """## What counts as a "buy" (the label)

The model predicts one yes/no event, not a price or a return size.

- A "buy" is set by a **triple-barrier** test (Lopez de Prado). From each day's close, draw three lines: an upper barrier at **+10%**, a lower barrier at **-5%**, and a time barrier **20 days** out.
- Walk forward day by day. If price reaches **+10% before** it falls to -5%, that day is a buy (**label = 1**).
- If it hits **-5% first**, or 20 days pass without reaching +10%, it is a **0**.
- On a day where both could have happened, we assume the **stop (-5%) hit first**, so the label never flatters itself.
- So the question the model answers is: will a +10% move arrive before a -5% drawdown within 20 days?

## How precision is scored

Every test day has two facts: what the model said (buy or not, at the 0.5 cut) and what actually happened. That gives four outcomes:

- **True positive** - said buy, and it was a buy. A good call.
- **False positive** - said buy, but it was not. A bad trade that spends real money.
- **False negative** - said no, but it was a buy. A missed chance.
- **True negative** - said no, and it was not. Correctly stood aside.

- **Precision = true positives / (true positives + false positives).** In plain words: every time the model shouts "buy", how often is it right?
- Precision only punishes bad trades (false positives) and ignores missed chances (false negatives). That is deliberate: a bad trade loses money now, a missed chance only costs an opportunity that comes around again.
- Recall is the mirror (of all the real buys, how many we caught). We rank models by precision, not recall, because being right when we act matters more than acting often.

## What the columns mean

- **date** - the day the evaluation was run.
- **evaluation type** - which kind of evaluation. *Head-to-head*: several models trained on the same train/test split and compared. Later types: walk-forward backtest, tuning sweep, stability check.
- **dataset** - rows and feature count used (e.g. 26,762r / 32f).
- **best model** - the model kept, chosen by the highest Best Model Precision.
- **test AUC** - area under the ROC curve on the out-of-sample test set: the chance the model scores a real buy above a non-buy. 0.50 = no skill (a coin flip), 1.00 = perfect ranking; it does not depend on a threshold.
- **Best Model Precision** - the precision of the chosen model: of the days it called a buy, the share that were genuine buys.
- **Always Buys Precision** - the precision a mindless model that calls every day a buy would score, i.e. the share of all test days that were buys. The baseline to beat.
- **Precision Change (%)** - how much better the best model is than always-buying: (Best Model Precision / Always Buys Precision - 1) x 100. +0% = no better than mindless; above 0 = adding value. Verify it from the two columns to its left.
- **Net P&L/trade** - Keller Metric 2 in one number: the average return of a trade the model takes (probability >= 0.60), after the 0.20% round-trip cost (Binance.com spot with BNB, plus slippage). Above 0 = a model-picked trade makes money after fees. This is the number to maximise when tuning.
- **trades** - how many trades the model would have taken in the test window (probability >= 0.60). Too few trades and the P&L is noise.
- **verdict** - GO only if Best Model Precision clearly beats Always Buys Precision and AUC clears 0.55; otherwise NO-GO.
- **record** - links to the full per-run report (markdown for the numbers, HTML for the charts).

Each per-run record also reports Keller Metric 1 (precision and recall at the 60/40 trading threshold) and Metric 3 (AUC split by volatility regime).
"""


def _b64(path):
    with open(path, "rb") as fh:
        return base64.b64encode(fh.read()).decode("ascii")


def _md_table(headers, rows):
    out = ["| " + " | ".join(headers) + " |",
           "| " + " | ".join("---" for _ in headers) + " |"]
    for r in rows:
        out.append("| " + " | ".join(str(c) for c in r) + " |")
    return "\n".join(out)


def _html_table(headers, rows):
    th = "".join(f"<th>{h}</th>" for h in headers)
    trs = "".join("<tr>" + "".join(f"<td>{c}</td>" for c in r) + "</tr>" for r in rows)
    return f"<table><thead><tr>{th}</tr></thead><tbody>{trs}</tbody></table>"


def _lift(prec, base):
    return prec / base if base else float("nan")


def _pct_change(prec, base):
    return (prec / base - 1.0) * 100.0 if base else float("nan")


# --------------------------------------------------------------------------- #
# Charts
# --------------------------------------------------------------------------- #
def _chart_compare(results, base_rate, path):
    names = [r["name"] for r in results]
    aucs = [r["auc"] for r in results]
    precs = [r["prec"] for r in results]
    x = np.arange(len(names)); w = 0.38
    fig, ax = plt.subplots(figsize=(7.5, 4.2))
    ax.bar(x - w / 2, aucs, w, label="test ROC-AUC", color=NAVY)
    ax.bar(x + w / 2, precs, w, label="precision (buy)", color=GREEN)
    ax.axhline(0.5, color=GREY, ls=":", lw=1)
    ax.axhline(base_rate, color=RED, ls="--", lw=1)
    ax.annotate("always-buys", xy=(0.995, base_rate), xycoords=ax.get_yaxis_transform(),
                ha="right", va="bottom", color=RED, fontsize=8)
    ax.set_xticks(x); ax.set_xticklabels(names)
    ax.set_ylim(0, max(0.6, max(aucs + precs) + 0.05))
    ax.set_title("Model comparison (out-of-sample test)")
    ax.legend(loc="upper left", framealpha=0.9)
    ax.grid(True, axis="y", ls=":", alpha=0.3)
    fig.tight_layout(); fig.savefig(path, dpi=160); plt.close(fig)


def _chart_roc(results, y_true, path):
    fig, ax = plt.subplots(figsize=(5.6, 5.2))
    for i, r in enumerate(results):
        fpr, tpr, _ = roc_curve(y_true, r["prob"])
        ax.plot(fpr, tpr, color=PALETTE[i % len(PALETTE)], lw=1.8,
                label=f"{r['name']} (AUC {r['auc']:.3f})")
    ax.plot([0, 1], [0, 1], color=GREY, ls="--", lw=1, label="coin flip")
    ax.set_xlabel("false positive rate"); ax.set_ylabel("true positive rate")
    ax.set_title("ROC curves"); ax.legend(loc="lower right", fontsize=8, framealpha=0.9)
    ax.grid(True, ls=":", alpha=0.3)
    fig.tight_layout(); fig.savefig(path, dpi=160); plt.close(fig)


def _chart_importance(fi_names, fi_values, chosen, path):
    order = np.argsort(fi_values)[::-1][:15]
    fig, ax = plt.subplots(figsize=(7.5, 5.2))
    ax.barh([fi_names[i] for i in order][::-1],
            [fi_values[i] for i in order][::-1], color=NAVY)
    ax.set_title(f"{chosen}: top 15 feature importances")
    fig.tight_layout(); fig.savefig(path, dpi=160); plt.close(fig)


def _chart_regime(regime_rows, path):
    labels = [r[0] for r in regime_rows]
    aucs = [float(r[3]) if r[3] != "n/a" else np.nan for r in regime_rows]
    fig, ax = plt.subplots(figsize=(6.0, 4.0))
    ax.bar(labels, aucs, color=NAVY)
    ax.axhline(0.5, color=GREY, ls="--", lw=1)
    ax.annotate("no skill", xy=(0.995, 0.5), xycoords=ax.get_yaxis_transform(),
                ha="right", va="bottom", color=GREY, fontsize=8)
    ax.set_ylim(0, max(0.6, np.nanmax(aucs) + 0.05) if np.isfinite(np.nanmax(aucs)) else 0.6)
    ax.set_ylabel("test ROC-AUC"); ax.set_title("Metric 3: AUC by volatility regime")
    ax.grid(True, axis="y", ls=":", alpha=0.3)
    fig.tight_layout(); fig.savefig(path, dpi=160); plt.close(fig)


def _chart_equity(eq, path):
    fig, ax = plt.subplots(figsize=(7.5, 4.0))
    ax.plot(range(1, len(eq) + 1), eq * 100.0, color=NAVY, lw=1.5)
    ax.axhline(0, color=GREY, ls="--", lw=1)
    ax.set_xlabel("trade number (test window, in date order)")
    ax.set_ylabel("cumulative net return (sum, 1 unit/trade) %")
    ax.set_title("Metric 2: equity curve (model trades, after costs)")
    ax.grid(True, ls=":", alpha=0.3)
    fig.tight_layout(); fig.savefig(path, dpi=160); plt.close(fig)


# --------------------------------------------------------------------------- #
# Metric 2: P&L after costs (Keller)
# --------------------------------------------------------------------------- #
def _pnl(prob, trade_ret, conf_hi, cost_frac):
    """Simulate the model's trades. A trade is taken on every test row where the
    probability clears conf_hi; its gross return is the realized triple-barrier
    return (trade_ret), and its net return subtracts the round-trip cost. Returns
    expectancy, win rate, the trade-by-trade equity curve, total return, a per-trade
    Sharpe, and max drawdown. Equal-weight, independent trades (sizing/overlap are
    the Chapter Two controls' concern)."""
    prob = np.asarray(prob)
    tr = np.asarray(trade_ret, dtype=float)
    mask = (prob >= conf_hi) & np.isfinite(tr)
    n = int(mask.sum())
    if n == 0:
        return dict(n=0, expectancy=float("nan"), winrate=float("nan"),
                    total=float("nan"), sharpe=float("nan"), maxdd=float("nan"), eq=None)
    nets = tr[mask] - cost_frac
    mean = float(nets.mean())
    sd = float(nets.std(ddof=1)) if n > 1 else 0.0
    eq = np.cumsum(nets)                      # additive, one unit per trade (honest for overlap)
    return dict(n=n, expectancy=mean, winrate=float((nets > 0).mean()), total=float(eq[-1]),
                sharpe=mean / sd if sd > 0 else float("nan"),
                tstat=mean / (sd / np.sqrt(n)) if sd > 0 else float("nan"), eq=eq)


# --------------------------------------------------------------------------- #
# Metric 3: regime-stratified AUC (Keller)
# --------------------------------------------------------------------------- #
def _regime_auc(y_true, prob, vol):
    """AUC of the chosen model within low / mid / high volatility terciles."""
    y = np.asarray(y_true); p = np.asarray(prob); v = np.asarray(vol, dtype=float)
    ok = np.isfinite(v)
    q33, q66 = np.nanquantile(v[ok], [1 / 3, 2 / 3])
    masks = [("low vol", ok & (v <= q33)),
             ("mid vol", ok & (v > q33) & (v < q66)),
             ("high vol", ok & (v >= q66))]
    rows = []
    for name, m in masks:
        n = int(m.sum())
        br = float(y[m].mean()) if n else float("nan")
        if n > 20 and len(np.unique(y[m])) == 2:
            auc = f"{roc_auc_score(y[m], p[m]):.3f}"
        else:
            auc = "n/a"
        rows.append([name, f"{n:,}", f"{br:.3f}", auc])
    return rows


# --------------------------------------------------------------------------- #
# Index
# --------------------------------------------------------------------------- #
def _read_existing_rows(md_path):
    rows = []
    if os.path.exists(md_path):
        sep_seen = False
        for ln in open(md_path).read().splitlines():
            s = ln.strip()
            if s.startswith("|---") or s.startswith("| ---"):
                sep_seen = True; continue
            if sep_seen and s.startswith("|"):
                rows.append([c.strip() for c in s.strip("|").split("|")])
    return rows


def _update_index(evals_dir, new_row):
    md_path = os.path.join(evals_dir, INDEX_FILE)
    rows = [[str(c) for c in new_row]] + _read_existing_rows(md_path)
    head = ("# Evaluation scores\n\n"
            "One row per evaluation run, newest first. Each row links to its full "
            "record.\n\n" + GLOSSARY + "\n## Runs\n\n")
    with open(md_path, "w") as fh:
        fh.write(head + _md_table(_INDEX_HEADERS, rows) + "\n")
    _write_index_pdf(evals_dir, rows)
    _write_index_docx(evals_dir, rows)
    return md_path


def _write_index_pdf(evals_dir, rows):
    """A portable PDF of the same scores table (drops the link-only 'record' column)."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import landscape, letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer,
                                         Table, TableStyle)
    except ImportError:
        print("(reportlab not installed: skipped evaluation-scores.pdf. pip install reportlab)")
        return None
    styles = getSampleStyleSheet()
    story = [Paragraph("Evaluation scores", styles["Title"]),
             Paragraph("One row per evaluation run, newest first.", styles["BodyText"]),
             Spacer(1, 10)]
    # the scores table first, so the numbers are the first thing you see
    keep = [i for i, h in enumerate(_INDEX_HEADERS) if h != "record"]
    data = [[_INDEX_HEADERS[i] for i in keep]] + \
           [[r[i] if i < len(r) else "" for i in keep] for r in rows]
    tbl = Table(data, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B3D66")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f4f6f8")]),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 16))
    # the definitions below the table, keeping their section headings for readability
    for line in GLOSSARY.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("## "):
            story.append(Spacer(1, 6))
            story.append(Paragraph(s[3:], styles["Heading3"]))
            continue
        txt = (s.lstrip("- ").replace("**", "").replace("*", "")
               .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
        story.append(Paragraph(txt, styles["BodyText"]))
    pdf_path = os.path.join(evals_dir, "evaluation-scores.pdf")
    SimpleDocTemplate(pdf_path, pagesize=landscape(letter), leftMargin=28,
                      rightMargin=28, topMargin=28, bottomMargin=28).build(story)
    return pdf_path


def _write_index_docx(evals_dir, rows):
    """A Word version that preserves the markdown report look: the scores table
    first, then the definitions as headings and bullets below it."""
    try:
        import docx
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.shared import Pt
    except ImportError:
        print("(python-docx not installed: skipped evaluation-scores.docx. pip install python-docx)")
        return None
    d = docx.Document()
    d.add_heading("Evaluation scores", level=0)
    d.add_paragraph("One row per evaluation run, newest first.")
    keep = [i for i, h in enumerate(_INDEX_HEADERS) if h != "record"]
    headers = [_INDEX_HEADERS[i] for i in keep]
    t = d.add_table(rows=1, cols=len(headers))
    try:
        t.style = "Light Grid Accent 1"
    except KeyError:
        t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        run = t.rows[0].cells[j].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8)
    for r in rows:
        cells = t.add_row().cells
        for j, i in enumerate(keep):
            run = cells[j].paragraphs[0].add_run(r[i] if i < len(r) else "")
            run.font.size = Pt(8)
    d.add_paragraph("")
    for line in GLOSSARY.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("## "):
            d.add_heading(s[3:], level=2)
        elif s.startswith("- "):
            d.add_paragraph(s[2:].replace("**", ""), style="List Bullet")
        else:
            d.add_paragraph(s.replace("**", ""))
    docx_path = os.path.join(evals_dir, "evaluation-scores.docx")
    d.save(docx_path)
    return docx_path


# --------------------------------------------------------------------------- #
# Head-to-head model-comparison record
# --------------------------------------------------------------------------- #
def write_comparison(evals_dir, results, y_true, meta):
    """Write the md + html + charts for one head-to-head model comparison and add a
    row to evaluation-scores.md.

    results: list of dicts {name, prob, cv_auc, auc, acc, prec, rec, conf{...}}.
    meta: dataset_rows, n_features, train_rows, test_rows, base_rate, cut, embargo,
          conf_hi, conf_lo, chosen, verdict, fi_names, fi_values (optional),
          regime_vol (optional 1d array for Metric 3), date (optional YYYYMMDD).
    """
    os.makedirs(evals_dir, exist_ok=True)
    cd = meta.get("date") or datetime.now(timezone.utc).strftime("%Y%m%d")
    hd = f"{cd[:4]}-{cd[4:6]}-{cd[6:]}"          # date folder, YYYY-MM-DD
    run_dir = os.path.join(evals_dir, hd)
    os.makedirs(run_dir, exist_ok=True)
    stem = f"eval-head-to-head-{cd}"
    y_true = np.asarray(y_true)
    base = meta["base_rate"]
    chosen = next(r for r in results if r["name"] == meta["chosen"])
    verdict = meta["verdict"]
    pctx = _pct_change(chosen["prec"], base)

    # charts
    _chart_compare(results, base, os.path.join(run_dir, f"{stem}-compare.png"))
    _chart_roc(results, y_true, os.path.join(run_dir, f"{stem}-roc.png"))
    charts = [("Model comparison", f"{stem}-compare.png"),
              ("ROC curves", f"{stem}-roc.png")]
    if meta.get("fi_names") is not None and meta.get("fi_values") is not None:
        _chart_importance(meta["fi_names"], meta["fi_values"], meta["chosen"],
                          os.path.join(run_dir, f"{stem}-importance.png"))
        charts.append((f"{meta['chosen']} feature importance", f"{stem}-importance.png"))

    # Metric 3: regime-stratified AUC of the chosen model
    regime_rows = None
    if meta.get("regime_vol") is not None:
        regime_rows = _regime_auc(y_true, chosen["prob"], meta["regime_vol"])
        _chart_regime(regime_rows, os.path.join(run_dir, f"{stem}-regime.png"))
        charts.append(("Metric 3: AUC by volatility regime", f"{stem}-regime.png"))

    # Metric 2: P&L after costs of the chosen model's trades
    pnl = None
    pnl_rows = None
    if meta.get("trade_ret") is not None:
        pnl = _pnl(chosen["prob"], meta["trade_ret"], meta["conf_hi"],
                   meta.get("cost_pct", 0.20) / 100.0)
        if pnl["n"]:
            _chart_equity(pnl["eq"], os.path.join(run_dir, f"{stem}-equity.png"))
            charts.append(("Metric 2: equity curve (after costs)", f"{stem}-equity.png"))
            pnl_rows = [
                [f"trades taken (prob >= {meta['conf_hi']:.2f})", f"{pnl['n']:,}"],
                ["win rate (net > 0)", f"{pnl['winrate']:.1%}"],
                ["net expectancy / trade", f"{pnl['expectancy'] * 100:+.2f}%"],
                ["per-trade Sharpe", f"{pnl['sharpe']:.3f}"],
                ["t-stat (expectancy vs 0; |t|>2 ~ significant)", f"{pnl['tstat']:+.2f}"],
            ]

    # tables
    ds_rows = [
        ["rows", f"{meta['dataset_rows']:,}"], ["features", meta["n_features"]],
        ["train rows", f"{meta['train_rows']:,}"], ["test rows", f"{meta['test_rows']:,}"],
        ["split date", meta["cut"]], ["embargo (days)", meta["embargo"]],
        ["always-buys precision (base rate)", f"{base:.3f}"],
        ["confidence band", f"{meta['conf_lo']} - {meta['conf_hi']}"],
    ]
    model_headers = ["model", "CV AUC", "test AUC", "precision", "recall",
                     "accuracy", "Precision Change (%)"]
    model_rows = [[r["name"], f"{r['cv_auc']:.3f}", f"{r['auc']:.3f}", f"{r['prec']:.3f}",
                   f"{r['rec']:.3f}", f"{r['acc']:.3f}", f"{_pct_change(r['prec'], base):+.1f}%"]
                  for r in results]
    conf_headers = ["model", "coverage", "precision", "recall", "F1"]
    conf_rows = [[r["name"], f"{r['conf']['coverage']:.0%}", f"{r['conf']['precision']:.3f}",
                  f"{r['conf']['recall']:.3f}", f"{r['conf']['f1']:.3f}"] for r in results]
    regime_headers = ["regime", "rows", "base rate", "test AUC"]

    # markdown
    md = [f"# Head-to-head model comparison ({hd})\n",
          f"**Verdict: {verdict}** - best model **{meta['chosen']}**, test AUC "
          f"{chosen['auc']:.3f}, best-model precision {chosen['prec']:.3f} vs always-buys "
          f"{base:.3f} (precision change {pctx:+.1f}%).\n",
          "## Dataset and split\n\n" + _md_table(["field", "value"], ds_rows) + "\n",
          "## Models, out-of-sample test\n\n"
          "Precision is TP/(TP+FP) at the 0.5 threshold; Precision Change (%) is (precision / always-buys - 1) x 100.\n\n"
          + _md_table(model_headers, model_rows) + "\n",
          "## Metric 1: confidence filter (act if p >= hi or p <= lo)\n\n"
          "Keller's rule: score only the high-conviction rows we would actually trade. "
          f"Read precision against always-buys precision ({base:.3f}).\n\n"
          + _md_table(conf_headers, conf_rows) + "\n"]
    if regime_rows is not None:
        md.append("## Metric 3: regime-stratified performance\n\n"
                  "AUC of the best model within low / mid / high volatility terciles "
                  "(by 30-day realized volatility). A model that only works in one regime "
                  "is fragile.\n\n" + _md_table(regime_headers, regime_rows) + "\n")
    if pnl_rows is not None:
        md.append("## Metric 2: simulated P&L (after costs)\n\n"
                  f"Trades = test days where the model's probability clears {meta['conf_hi']:.2f}, "
                  "each held under the +10% / -5% / 20-day triple barrier, minus a "
                  f"{meta.get('cost_pct', 0.20):.2f}% round-trip cost (Binance.com spot with BNB, "
                  "plus slippage). Equal-weight, independent trades; position sizing and overlap "
                  "are the Chapter Two controls' job, not modelled here. The per-trade Sharpe and "
                  "t-stat say whether the expectancy is signal or noise: |t| above about 2 is the "
                  "rough bar for significance. The equity curve is the cumulative sum of net trade "
                  "returns at one unit per trade, not a sized portfolio.\n\n"
                  + _md_table(["metric", "value"], pnl_rows) + "\n")
    else:
        md.append("## Metric 2: simulated P&L\n\nNo trades cleared the confidence cut, or trade "
                  "returns were not supplied to the writer.\n")
    md.append("## Charts\n")
    for title, fn in charts:
        md.append(f"\n**{title}**\n\n![{title}]({fn})\n")
    md_path = os.path.join(run_dir, f"{stem}.md")
    with open(md_path, "w") as fh:
        fh.write("\n".join(md) + "\n")

    # html (charts embedded base64, self-contained)
    imgs = "".join(
        f"<h3>{title}</h3><img src='data:image/png;base64,"
        f"{_b64(os.path.join(run_dir, fn))}' style='max-width:760px;width:100%'>"
        for title, fn in charts)
    css = ("body{font-family:-apple-system,Segoe UI,Roboto,sans-serif;max-width:860px;"
           "margin:2rem auto;color:#222;padding:0 1rem}h1{color:#0B3D66}"
           "h2{color:#0B3D66;border-bottom:1px solid #eee;padding-bottom:4px}"
           "table{border-collapse:collapse;margin:0.5rem 0}"
           "th,td{border:1px solid #ddd;padding:5px 10px;text-align:right;font-size:14px}"
           "th:first-child,td:first-child{text-align:left}th{background:#f4f6f8}.v{font-size:16px}")
    regime_html = ("<h2>Metric 3: regime-stratified performance</h2>"
                   + _html_table(regime_headers, regime_rows)) if regime_rows is not None else ""
    metric2_html = ("<h2>Metric 2: simulated P&L (after costs)</h2>"
                    + _html_table(["metric", "value"], pnl_rows)) if pnl_rows is not None else ""
    html = (f"<!doctype html><html><head><meta charset='utf-8'><title>{stem}</title>"
            f"<style>{css}</style></head><body>"
            f"<h1>Head-to-head model comparison ({hd})</h1>"
            f"<p class='v'><b>Verdict: {verdict}</b> - best model <b>{meta['chosen']}</b>, "
            f"test AUC {chosen['auc']:.3f}, best-model precision {chosen['prec']:.3f} vs "
            f"always-buys {base:.3f} (precision change {pctx:+.1f}%).</p>"
            f"<h2>Dataset and split</h2>{_html_table(['field', 'value'], ds_rows)}"
            f"<h2>Models, out-of-sample test</h2>{_html_table(model_headers, model_rows)}"
            f"<h2>Metric 1: confidence filter</h2>{_html_table(conf_headers, conf_rows)}"
            f"{metric2_html}{regime_html}<h2>Charts</h2>{imgs}</body></html>")
    html_path = os.path.join(run_dir, f"{stem}.html")
    with open(html_path, "w") as fh:
        fh.write(html)

    # index row (links point into the date folder)
    record = f"[md]({hd}/{stem}.md) / [html]({hd}/{stem}.html)"
    pnl_cell = f"{pnl['expectancy'] * 100:+.2f}%" if (pnl and pnl["n"]) else "n/a"
    trades_cell = f"{pnl['n']:,}" if (pnl and pnl["n"]) else "0"
    _update_index(evals_dir, [hd, "head-to-head",
                              f"{meta['dataset_rows']:,}r / {meta['n_features']}f",
                              meta["chosen"], f"{chosen['auc']:.3f}", f"{chosen['prec']:.3f}",
                              f"{base:.3f}", f"{pctx:+.1f}%", pnl_cell, trades_cell,
                              verdict, record])
    return {"md": md_path, "html": html_path, "charts": charts}
