"""Render trader-workflow.ipynb to docx with the operator's formatting rules.

One command replaces the fragile render-then-fix cycle:

    .venv/bin/python inputs/render_workflow_docx.py

Pipeline: pandoc renders the notebook against reference.docx (which carries the
page header with the repo URL, the 8pt + grey-shaded Source Code style, and all
other operator styles), then a post-pass stamps 8pt directly on every code and
table run, because pandoc emits its own 9pt direct run sizes that would
override the style. Never render the docx without this script: a bare pandoc
run regresses the operator's formatting.
"""

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

REPO = Path(__file__).resolve().parents[1]
NB = REPO / "02-runtime" / "trader-workflow.ipynb"
OUT = REPO / "02-runtime" / "trader-workflow.docx"
REF = REPO / "02-runtime" / "reference.docx"


def main() -> int:
    subprocess.run(
        ["/usr/local/bin/pandoc", str(NB), "-o", str(OUT),
         f"--reference-doc={REF}", "--toc", f"--resource-path={NB.parent}"],
        check=True,
    )

    doc = Document(str(OUT))
    code = table = 0
    for p in doc.paragraphs:
        if p.style.name == "Source Code":
            for r in p.runs:
                r.font.size = Pt(8)
                code += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)
                        table += 1
    doc.save(str(OUT))

    hdr = " ".join(p.text for p in doc.sections[0].header.paragraphs).strip()
    sc = doc.styles["Source Code"]
    print(f"rendered {OUT.name}: {code} code runs + {table} table runs at 8pt; "
          f"header='{hdr}'; SourceCode {sc.font.size.pt if sc.font.size else '?'}pt")
    return 0


if __name__ == "__main__":
    sys.exit(main())
